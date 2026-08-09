from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parent
TEMPLATE = ROOT / "proposta-tecnica-de-la-prova-OFICIAL.docx"
SOURCE = ROOT / "proposta-pulsepath-contingut.md"
OUTPUT = ROOT / "Proposta_tecnica_PulsePath_Ajuntament_Barcelona_FINAL.docx"


HEADING_MAP = {
    "Descripció": "Descripció",
    "Objectius de la prova": "Objectius de la prova",
    "Antecedents": "Antecedents",
    "Nivell de maduresa tecnològica": "Nivell de maduresa tecnològica",
    "El valor de provar en un entorn real": "El valor de provar en un entorn real",
    "Coherència amb els objectius estratègics municipals": "Coherència amb els objectius estratègics municipals",
    "Durada prevista": "Durada prevista",
    "Què ens fa diferents? (innovació i valor diferencial)": "Què us fa diferents? (innovació i valor diferencial)",
    "Encaix amb Barcelona": "Encaix amb Barcelona",
    "Potencial de creixement": "Potencial de creixement",
    "Metodologia": "Metodologia",
    "Cronograma i fases d'execució": "Cronograma i fases d’execució",
    "Actius necessaris": "Actius necessaris",
    "Necessitats tècniques i adequació": "Necessitats tècniques i adequació",
    "Capacitat d'execució": "Capacitat d’execució",
    "Integració urbana": "Integració urbana",
    "Seguretat i prevenció": "Seguretat i prevenció",
    "Compromís amb la diversitat i inclusió": "Compromís amb la diversitat i inclusió",
    "Criteris socioambientals i d'impacte de gènere": "Criteris socioambientals i d’impacte de gènere",
    "Valor públic i retorn social": "Valor públic i retorn social",
    "Aprenentatge normatiu i millora de la gestió": "Aprenentatge normatiu i millora de la gestió",
    "Indicadors clau (KPI)": "Indicadors clau (KPI)",
}


def normalize(text: str) -> str:
    return " ".join(text.replace("’", "'").split()).strip().lower()


def parse_h3_sections(markdown: str) -> dict[str, list[str]]:
    sections: dict[str, list[str]] = {}
    current: str | None = None
    for line in markdown.splitlines():
        if line.startswith("### "):
            current = line[4:].strip()
            sections[current] = []
        elif line.startswith("## "):
            current = None
        elif current is not None:
            sections[current].append(line.rstrip())
    return sections


def find_paragraph(doc: Document, text: str):
    target = normalize(text)
    for paragraph in doc.paragraphs:
        if normalize(paragraph.text) == target:
            return paragraph
    raise ValueError(f"No s'ha trobat el títol de secció: {text}")


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def next_paragraph(paragraph):
    node = paragraph._p.getnext()
    while node is not None:
        if node.tag.endswith("}p"):
            from docx.text.paragraph import Paragraph

            return Paragraph(node, paragraph._parent)
        node = node.getnext()
    return None


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill", fill)
    tc_pr.append(shd)


def add_inline_markdown(paragraph, text: str) -> None:
    text = text.strip()
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def add_paragraph_after(doc: Document, anchor_xml, text: str, style: str = "Normal"):
    paragraph = doc.add_paragraph(style=style)
    add_inline_markdown(paragraph, text)
    anchor_xml.addnext(paragraph._p)
    return paragraph._p


def add_table_after(doc: Document, anchor_xml, rows: list[list[str]]):
    table = doc.add_table(rows=1, cols=len(rows[0]))
    if "Table Grid" in [style.name for style in doc.styles]:
        table.style = "Table Grid"
    table.autofit = True

    for index, value in enumerate(rows[0]):
        cell = table.rows[0].cells[index]
        cell.text = value
        set_cell_shading(cell, "D9EAF7")
        for run in cell.paragraphs[0].runs:
            run.bold = True

    for row_values in rows[1:]:
        cells = table.add_row().cells
        for index, value in enumerate(row_values):
            cells[index].text = value

    anchor_xml.addnext(table._tbl)
    return table._tbl


def split_blocks(lines: list[str]):
    blocks: list[tuple[str, object]] = []
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue

        if line.startswith("|") and i + 1 < len(lines) and re.match(r"^\|?[\s:|-]+\|?$", lines[i + 1].strip()):
            table_lines = [line]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            rows = [
                [cell.strip() for cell in table_line.strip("|").split("|")]
                for table_line in table_lines
            ]
            blocks.append(("table", rows))
            continue

        if re.match(r"^[-*]\s+", line):
            blocks.append(("bullet", "• " + re.sub(r"^[-*]\s+", "", line)))
            i += 1
            continue

        if re.match(r"^\d+\.\s+", line):
            blocks.append(("number", line))
            i += 1
            continue

        paragraph_lines = [line]
        i += 1
        while i < len(lines):
            candidate = lines[i].strip()
            if (
                not candidate
                or candidate.startswith("|")
                or re.match(r"^[-*]\s+", candidate)
                or re.match(r"^\d+\.\s+", candidate)
            ):
                break
            paragraph_lines.append(candidate)
            i += 1
        blocks.append(("paragraph", " ".join(paragraph_lines)))
    return blocks


def insert_section_content(doc: Document, heading, lines: list[str]) -> None:
    instruction = next_paragraph(heading)
    if instruction is not None:
        remove_paragraph(instruction)

    anchor = heading._p
    for kind, value in split_blocks(lines):
        if kind == "table":
            anchor = add_table_after(doc, anchor, value)
        elif kind == "bullet":
            anchor = add_paragraph_after(doc, anchor, value)
        elif kind == "number":
            anchor = add_paragraph_after(doc, anchor, value)
        else:
            anchor = add_paragraph_after(doc, anchor, value)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Source Sans Pro"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    if "Caption" not in [style.name for style in doc.styles]:
        doc.styles.add_style("Caption", WD_STYLE_TYPE.PARAGRAPH)
    caption = doc.styles["Caption"]
    caption.font.name = "Source Sans Pro"
    caption.font.size = Pt(9)


def append_references_and_evidence(doc: Document) -> None:
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)

    heading = doc.add_paragraph("Referències públiques", style="Heading 1")
    references = [
        (
            "Ajuntament de Barcelona. II Pla de Salut Mental de Barcelona 2023-2030.",
            "https://ajuntament.barcelona.cat/sanitatisalut/ca/canal/pla-de-salut-mental",
        ),
        (
            "Ajuntament de Barcelona. Acord de ciutat per cuidar la salut mental a la feina.",
            "https://ajuntament.barcelona.cat/sanitatisalut/ca/canal/acord-de-ciutat-cuidar-la-salut-mental-la-feina",
        ),
        (
            "Ajuntament de Barcelona. Pla per a la reducció de l'absentisme del personal municipal (2026).",
            "https://ajuntament.barcelona.cat/transparencia/",
        ),
        (
            "Ajuntament de Barcelona. Agenda 2030 de Barcelona.",
            "https://ajuntament.barcelona.cat/agenda2030/",
        ),
    ]
    for label, url in references:
        paragraph = doc.add_paragraph()
        paragraph.add_run(f"• {label} {url}")

    note = doc.add_paragraph()
    run = note.add_run(
        "La referència al paradigma PVT i a publicacions de Mathias Basner no implica "
        "una validació de la implementació de PulsePath ni un aval formal."
    )
    run.italic = True

    doc.add_paragraph("Evidència visual del prototip", style="Heading 1")

    images = [
        (
            ROOT / "evidencia-app-local.png",
            "Figura 1. Accés de la persona participant mitjançant un codi pseudònim. Captura del prototip en entorn local.",
            Inches(6.3),
        ),
        (
            ROOT / "evidencia-app-consentimiento.png",
            "Figura 2. Consentiment informat abans de començar. PulsePath s'identifica com a eina no clínica.",
            Inches(6.3),
        ),
        (
            ROOT / "evidencia-dashboard-portada.png",
            "Figura 3. Quadre de comandament agregat. Les dades de la captura són sintètiques i només serveixen de demostració.",
            Inches(6.3),
        ),
    ]
    for image_path, caption, width in images:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        paragraph.add_run().add_picture(str(image_path), width=width)
        caption_paragraph = doc.add_paragraph(caption, style="Caption")
        caption_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def main() -> None:
    markdown = SOURCE.read_text(encoding="utf-8")
    sections = parse_h3_sections(markdown)
    doc = Document(TEMPLATE)
    configure_styles(doc)

    title = markdown.splitlines()[0].removeprefix("# ").strip()
    title_paragraph = doc.paragraphs[4]
    title_paragraph.clear()
    title_run = title_paragraph.add_run(title)
    title_run.italic = False
    title_run.font.name = "Source Sans Pro"
    title_run.font.size = Pt(24)
    title_run.bold = True

    cover_meta = doc.paragraphs[14]
    cover_meta.style = doc.styles["Normal"]
    cover_meta.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    cover_meta.add_run("Promotor: Max Borra Palau · PulsePath\n").bold = True
    cover_meta.add_run("Barcelona · Juliol de 2026\n")
    cover_meta.add_run("Tràmit 20260001771 · Accés als Espais d'experimentació")

    for source_heading, template_heading in HEADING_MAP.items():
        if source_heading not in sections:
            raise ValueError(f"Falta contingut per a: {source_heading}")
        heading = find_paragraph(doc, template_heading)
        insert_section_content(doc, heading, sections[source_heading])

    append_references_and_evidence(doc)

    doc.core_properties.title = title
    doc.core_properties.subject = "Proposta tècnica de prova als Espais d'experimentació de Barcelona"
    doc.core_properties.author = "Max Borra Palau · PulsePath"
    doc.core_properties.keywords = "PulsePath, Barcelona, fatiga, estrès, benestar laboral, pilot"

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
