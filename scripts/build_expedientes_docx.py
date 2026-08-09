"""Genera els DOCX dels expedients de Barcelona i València a partir del markdown.

Sortides:
  Barcelona (plantilla oficial de l'Ajuntament):
    docs/ayuntamiento-espais-experimentacio/Proposta_tecnica_PulsePath_Ajuntament_Barcelona_FINAL_v3.docx
    docs/ayuntamiento-espais-experimentacio/declaracions/Declaracions_responsables_PulsePath.docx
  València (documents autònoms):
    docs/ayuntamiento-valencia/1_Entrevista_Informativa_PulsePath.docx
    docs/ayuntamiento-valencia/2_Cuestionario_Inicial_PulsePath.docx
    docs/ayuntamiento-valencia/3_Protocolo_Pruebas_PulsePath.docx

Ús:  python scripts/build_expedientes_docx.py
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
BCN = ROOT / "docs" / "ayuntamiento-espais-experimentacio"
DECL = BCN / "declaracions"
VLC = ROOT / "docs" / "ayuntamiento-valencia"

TEMPLATE = BCN / "proposta-tecnica-de-la-prova-OFICIAL.docx"
SOURCE = BCN / "proposta-pulsepath-contingut.md"
OUTPUT_BCN = BCN / "Proposta_tecnica_PulsePath_Ajuntament_Barcelona_FINAL_v3.docx"
OUTPUT_DECL = DECL / "Declaracions_responsables_PulsePath.docx"

BODY_FONT = "Source Sans Pro"
MONO_FONT = "Consolas"
HEADER_FILL = "D9EAF7"
MISSING_ASSET_COLOR = RGBColor(0xB1, 0x4B, 0x00)

# Seccions H3 del markdown -> títols fixos de la plantilla oficial.
HEADING_MAP = {
    "Descripció": "Descripció",
    "Objectius de la prova": "Objectius de la prova",
    "Antecedents": "Antecedents",
    "Nivell de maduresa tecnològica": "Nivell de maduresa tecnològica",
    "El valor de provar en un entorn real": "El valor de provar en un entorn real",
    "Coherència amb els objectius estratègics municipals": "Coherència amb els objectius estratègics municipals",
    "Durada prevista": "Durada prevista",
    "Què ens fa diferents? (innovació i valor diferencial)": "Què us fa diferents? (innovació i valor diferencial)",
    "Encaix amb Barcelona": "Encaix amb Barcelona",
    "Potencial de creixement": "Potencial de creixement",
    "Metodologia": "Metodologia",
    "Cronograma i fases d'execució": "Cronograma i fases d’execució",
    "Actius necessaris": "Actius necessaris",
    "Necessitats tècniques i adequació": "Necessitats tècniques i adequació",
    "Capacitat d'execució": "Capacitat d’execució",
    "Integració urbana": "Integració urbana",
    "Seguretat i prevenció": "Seguretat i prevenció",
    "Compromís amb la diversitat i inclusió": "Compromís amb la diversitat i inclusió",
    "Criteris socioambientals i d'impacte de gènere": "Criteris socioambientals i d’impacte de gènere",
    "Valor públic i retorn social": "Valor públic i retorn social",
    "Aprenentatge normatiu i millora de la gestió": "Aprenentatge normatiu i millora de la gestió",
    "Indicadors clau (KPI)": "Indicadors clau (KPI)",
}

# La plantilla oficial no té ranura per a "Compliment algorítmic" (art. 10.3.b.6).
# Es fusiona dins del criteri d'elegibilitat del qual depèn (art. 12.2.a).
MERGE_INTO = {"Compliment algorítmic": "Nivell de maduresa tecnològica"}


# --------------------------------------------------------------------------
# Markdown → blocs
# --------------------------------------------------------------------------

HTML_COMMENT_RE = re.compile(r"<!--.*?-->", re.DOTALL)


def strip_comments(markdown: str) -> str:
    """Les notes internes HTML no han d'aparèixer en un document per signar."""
    return HTML_COMMENT_RE.sub("", markdown)


def parse_h3_sections(markdown: str) -> dict[str, list[str]]:
    sections: dict[str, list[str]] = {}
    current: str | None = None
    for line in markdown.splitlines():
        if line.startswith("### "):
            current = line[4:].strip()
            sections[current] = []
        elif line.startswith("## "):
            current = None
        elif current is not None:
            sections[current].append(line.rstrip())
    return sections


def split_blocks(lines: list[str]) -> list[tuple[str, object]]:
    """Converteix línies de markdown en blocs (tipus, valor)."""
    blocks: list[tuple[str, object]] = []
    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.strip()

        if not line or line in {"---", "***", "___"}:
            i += 1
            continue

        if line.startswith("#"):
            level = len(line) - len(line.lstrip("#"))
            blocks.append((f"h{level}", line[level:].strip()))
            i += 1
            continue

        # Taula: fila de capçalera seguida de la fila separadora.
        if line.startswith("|") and i + 1 < len(lines) and re.match(r"^\|?[\s:|-]+\|?$", lines[i + 1].strip()):
            table_lines = [line]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            rows = [[c.strip() for c in tl.strip("|").split("|")] for tl in table_lines]
            blocks.append(("table", rows))
            continue

        if re.match(r"^[-*]\s+", line):
            indent = len(raw) - len(raw.lstrip())
            blocks.append(("bullet2" if indent >= 2 else "bullet", re.sub(r"^[-*]\s+", "", line)))
            i += 1
            continue

        if re.match(r"^\d+\.\s+", line):
            indent = len(raw) - len(raw.lstrip())
            blocks.append(("number2" if indent >= 2 else "number", line))
            i += 1
            continue

        if set(line) == {"_"}:
            blocks.append(("rule", line))
            i += 1
            continue

        paragraph_lines = [line]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if (
                not nxt
                or nxt.startswith("|")
                or nxt.startswith("#")
                or re.match(r"^[-*]\s+", nxt)
                or re.match(r"^\d+\.\s+", nxt)
                or set(nxt) == {"_"}
                or nxt in {"---", "***", "___"}
            ):
                break
            paragraph_lines.append(nxt)
            i += 1
        blocks.append(("paragraph", " ".join(paragraph_lines)))
    return blocks


# --------------------------------------------------------------------------
# Renderitzat inline
# --------------------------------------------------------------------------

INLINE_RE = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|`.+?`)")


def add_inline_markdown(paragraph, text: str) -> None:
    for part in INLINE_RE.split(text.strip()):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            paragraph.add_run(part[1:-1]).italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = MONO_FONT
            run.font.size = Pt(9.5)
        else:
            paragraph.add_run(part)


def set_cell_shading(cell, fill: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shd)


# --------------------------------------------------------------------------
# Renderitzat de blocs (mode "append", per a documents autònoms)
# --------------------------------------------------------------------------


def render_blocks(doc: Document, blocks, heading_offset: int = 0) -> None:
    for kind, value in blocks:
        if kind.startswith("h"):
            level = min(int(kind[1:]) + heading_offset, 4)
            doc.add_paragraph(str(value), style=f"Heading {max(level, 1)}")
        elif kind == "table":
            render_table(doc, value)
        elif kind in {"bullet", "bullet2"}:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5 if kind == "bullet2" else 0.25)
            p.add_run("◦ " if kind == "bullet2" else "• ")
            add_inline_markdown(p, str(value))
        elif kind in {"number", "number2"}:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5 if kind == "number2" else 0.25)
            add_inline_markdown(p, str(value))
        elif kind == "rule":
            doc.add_paragraph(str(value))
        else:
            add_inline_markdown(doc.add_paragraph(), str(value))


def render_table(doc: Document, rows: list[list[str]]):
    table = doc.add_table(rows=1, cols=len(rows[0]))
    if "Table Grid" in [s.name for s in doc.styles]:
        table.style = "Table Grid"
    table.autofit = True
    for idx, value in enumerate(rows[0]):
        cell = table.rows[0].cells[idx]
        cell.text = ""
        add_inline_markdown(cell.paragraphs[0], value)
        set_cell_shading(cell, HEADER_FILL)
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for row_values in rows[1:]:
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            if idx >= len(cells):
                break
            cells[idx].text = ""
            add_inline_markdown(cells[idx].paragraphs[0], value)
    return table


# --------------------------------------------------------------------------
# Renderitzat de blocs (mode "insert after", per a la plantilla oficial)
# --------------------------------------------------------------------------


def next_paragraph(paragraph):
    from docx.text.paragraph import Paragraph

    node = paragraph._p.getnext()
    while node is not None:
        if node.tag.endswith("}p"):
            return Paragraph(node, paragraph._parent)
        node = node.getnext()
    return None


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def find_paragraph(doc: Document, text: str):
    target = " ".join(text.replace("’", "'").split()).strip().lower()
    for paragraph in doc.paragraphs:
        if " ".join(paragraph.text.replace("’", "'").split()).strip().lower() == target:
            return paragraph
    raise ValueError(f"No s'ha trobat el títol de secció a la plantilla: {text}")


def insert_section_content(doc: Document, heading, lines: list[str]) -> None:
    """Substitueix el text d'instruccions de la plantilla pel contingut real."""
    instruction = next_paragraph(heading)
    if instruction is not None:
        remove_paragraph(instruction)

    anchor = heading._p
    for kind, value in split_blocks(lines):
        if kind == "table":
            table = render_table(doc, value)
            anchor.addnext(table._tbl)
            anchor = table._tbl
            continue

        p = doc.add_paragraph()
        if kind.startswith("h"):
            # La plantilla té títols fixos; els subtítols del markdown es
            # representen en negreta per no trencar-ne la jerarquia.
            run = p.add_run(str(value))
            run.bold = True
            run.font.size = Pt(11.5)
            p.paragraph_format.space_before = Pt(10)
        elif kind in {"bullet", "bullet2"}:
            p.paragraph_format.left_indent = Inches(0.5 if kind == "bullet2" else 0.25)
            p.add_run("◦ " if kind == "bullet2" else "• ")
            add_inline_markdown(p, str(value))
        elif kind in {"number", "number2"}:
            p.paragraph_format.left_indent = Inches(0.5 if kind == "number2" else 0.25)
            add_inline_markdown(p, str(value))
        else:
            add_inline_markdown(p, str(value))

        anchor.addnext(p._p)
        anchor = p._p


# --------------------------------------------------------------------------
# Estils
# --------------------------------------------------------------------------


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    if "Caption" not in [s.name for s in doc.styles]:
        doc.styles.add_style("Caption", WD_STYLE_TYPE.PARAGRAPH)
    caption = doc.styles["Caption"]
    caption.font.name = BODY_FONT
    caption.font.size = Pt(9)
    caption.font.italic = True


def set_core_properties(doc: Document, title: str, subject: str, keywords: str) -> None:
    doc.core_properties.title = title
    doc.core_properties.subject = subject
    doc.core_properties.author = "Max Borra Palau · PulsePath"
    doc.core_properties.keywords = keywords


# --------------------------------------------------------------------------
# Barcelona
# --------------------------------------------------------------------------


def append_references_and_evidence(doc: Document) -> None:
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    doc.add_paragraph("Referències públiques", style="Heading 1")

    references = [
        (
            "Reglament (UE) 2024/1689 (AI Act), en la redacció resultant del Reglament (UE) 2026/1744.",
            "https://eur-lex.europa.eu/eli/reg/2024/1689/oj",
        ),
        (
            "Reglament (UE) 2026/1744 («Digital Omnibus» sobre IA), DOUE de 24 de juliol de 2026.",
            "https://eur-lex.europa.eu/legal-content/CA/TXT/?uri=OJ:L_202601744",
        ),
        (
            "Ajuntament de Barcelona. II Pla de Salut Mental de Barcelona 2023-2030.",
            "https://ajuntament.barcelona.cat/sanitatisalut/ca/canal/pla-de-salut-mental",
        ),
        (
            "Ajuntament de Barcelona. Acord de ciutat per cuidar la salut mental a la feina.",
            "https://ajuntament.barcelona.cat/sanitatisalut/ca/canal/acord-de-ciutat-cuidar-la-salut-mental-la-feina",
        ),
        (
            "Ajuntament de Barcelona. Pla per a la reducció de l'absentisme del personal municipal (2026).",
            "https://ajuntament.barcelona.cat/transparencia/",
        ),
        (
            "Ajuntament de Barcelona. Agenda 2030 de Barcelona.",
            "https://ajuntament.barcelona.cat/agenda2030/",
        ),
    ]
    for label, url in references:
        doc.add_paragraph().add_run(f"• {label} {url}")

    note = doc.add_paragraph()
    note.add_run(
        "La referència al paradigma PVT i a publicacions de Mathias Basner no implica "
        "una validació de la implementació de PulsePath ni un aval formal."
    ).italic = True

    doc.add_paragraph("Evidència visual del prototip", style="Heading 1")

    disclaimer = doc.add_paragraph()
    disclaimer.add_run(
        "Nota: les captures següents corresponen a l'entorn de proves i acrediten l'estat "
        "actual del sistema. No hi ha una URL pública de demostració desplegada; es pot "
        "concertar una demostració sota petició."
    ).italic = True

    images = [
        (
            BCN / "evidencia-app-local.png",
            "Figura 1. Accés de la persona participant mitjançant un codi pseudònim. Captura de l'entorn de proves.",
        ),
        (
            BCN / "evidencia-app-consentimiento.png",
            "Figura 2. Consentiment informat abans de començar. PulsePath s'identifica com a eina no clínica.",
        ),
        (
            BCN / "evidencia-dashboard-portada.png",
            "Figura 3. Quadre de comandament agregat amb supressió K ≥ 5. Les dades de la captura són sintètiques.",
        ),
    ]
    for image_path, caption in images:
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        if image_path.exists():
            p.add_run().add_picture(str(image_path), width=Inches(6.3))
        else:
            placeholder = p.add_run(f"[PENDENT DE CAPTURA — {image_path.name}]")
            placeholder.bold = True
            placeholder.font.color.rgb = MISSING_ASSET_COLOR
        cap = doc.add_paragraph(caption, style="Caption")
        cap.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def parse_h2_section(markdown: str, prefix: str) -> tuple[str, list[str]] | None:
    """Retorna (títol, línies) de la primera secció H2 que comenci per `prefix`."""
    title: str | None = None
    lines: list[str] = []
    for line in markdown.splitlines():
        if line.startswith("## "):
            if title is not None:
                break
            heading = line[3:].strip()
            if heading.startswith(prefix):
                title = heading
            continue
        if title is not None:
            lines.append(line.rstrip())
    return (title, lines) if title else None


def append_annexes(doc: Document, markdown: str) -> None:
    for prefix in ("Annex A",):
        parsed = parse_h2_section(markdown, prefix)
        if not parsed:
            continue
        title, lines = parsed
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        doc.add_paragraph(title, style="Heading 1")
        render_blocks(doc, split_blocks(lines))


def build_barcelona() -> Path:
    markdown = strip_comments(SOURCE.read_text(encoding="utf-8"))
    sections = parse_h3_sections(markdown)

    # Fusió de les seccions sense ranura pròpia a la plantilla oficial.
    for source_heading, target_heading in MERGE_INTO.items():
        if source_heading in sections:
            sections[target_heading] = (
                sections[target_heading] + ["", f"#### {source_heading}", ""] + sections.pop(source_heading)
            )

    doc = Document(TEMPLATE)
    configure_styles(doc)

    title = markdown.splitlines()[0].removeprefix("# ").strip()
    title_paragraph = doc.paragraphs[4]
    title_paragraph.clear()
    run = title_paragraph.add_run(title)
    run.bold = True
    run.italic = False
    run.font.name = BODY_FONT
    run.font.size = Pt(24)

    cover = doc.paragraphs[14]
    cover.style = doc.styles["Normal"]
    cover.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    cover.add_run("Promotor: Max Borra Palau · PulsePath\n").bold = True
    cover.add_run("Barcelona · 9 d'agost de 2026\n")
    cover.add_run("Tràmit 20260001771 · Accés als Espais d'experimentació")

    for source_heading, template_heading in HEADING_MAP.items():
        if source_heading not in sections:
            raise ValueError(f"Falta contingut per a la secció: {source_heading}")
        insert_section_content(doc, find_paragraph(doc, template_heading), sections[source_heading])

    append_references_and_evidence(doc)
    append_annexes(doc, markdown)

    set_core_properties(
        doc,
        title,
        "Proposta tècnica de prova als Espais d'experimentació de Barcelona",
        "PulsePath, Barcelona, fatiga, estrès, benestar laboral, pilot, AI Act",
    )
    doc.save(OUTPUT_BCN)
    return OUTPUT_BCN


def build_declaracions() -> Path:
    files = sorted(p for p in DECL.glob("0*.md"))
    if len(files) != 6:
        raise ValueError(f"S'esperaven 6 declaracions, se n'han trobat {len(files)}: {[f.name for f in files]}")

    doc = Document()
    configure_styles(doc)

    doc.add_paragraph("Declaracions responsables", style="Title")
    subtitle = doc.add_paragraph()
    subtitle.add_run(
        "PulsePath · Accés als Espais d'experimentació de Barcelona · Tràmit 20260001771\n"
    ).bold = True
    subtitle.add_run("Article 10.3.b de l'Ordenança dels espais d'experimentació de Barcelona")

    intro = doc.add_paragraph()
    intro.add_run(
        "Aquest document recull les sis declaracions responsables exigides per l'article 10.3.b. "
        "Cada declaració es pot presentar també com a document independent. "
        "Cal emplenar a mà les dades personals del declarant i signar abans de presentar."
    ).italic = True

    for index, path in enumerate(files):
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        blocks = split_blocks(strip_comments(path.read_text(encoding="utf-8")).splitlines())
        render_blocks(doc, blocks)
        if index < len(files) - 1:
            doc.add_paragraph()

    set_core_properties(
        doc,
        "Declaracions responsables — PulsePath",
        "Declaracions responsables de l'article 10.3.b de l'Ordenança",
        "PulsePath, Barcelona, declaracions responsables, AI Act, RGPD",
    )
    doc.save(OUTPUT_DECL)
    return OUTPUT_DECL


# --------------------------------------------------------------------------
# València
# --------------------------------------------------------------------------

VALENCIA_DOCS = [
    ("1-entrevista-informativa.md", "1_Entrevista_Informativa_PulsePath.docx", "Entrevista informativa"),
    ("2-cuestionario-inicial.md", "2_Cuestionario_Inicial_PulsePath.docx", "Cuestionario inicial"),
    ("3-protocolo-de-pruebas.md", "3_Protocolo_Pruebas_PulsePath.docx", "Protocolo de pruebas"),
]


def build_valencia() -> list[Path]:
    outputs: list[Path] = []
    for source_name, output_name, subject in VALENCIA_DOCS:
        source = VLC / source_name
        markdown = strip_comments(source.read_text(encoding="utf-8"))
        blocks = split_blocks(markdown.splitlines())

        doc = Document()
        configure_styles(doc)
        render_blocks(doc, blocks)

        title = markdown.splitlines()[0].removeprefix("# ").strip()
        set_core_properties(
            doc,
            f"{title} — PulsePath",
            f"Sandbox Urbano de València · {subject}",
            "PulsePath, València, Sandbox Urbano, fatiga, bienestar laboral",
        )
        output = VLC / output_name
        doc.save(output)
        outputs.append(output)
    return outputs


def main() -> None:
    generated = [build_barcelona(), build_declaracions(), *build_valencia()]
    for path in generated:
        print(f"{path.relative_to(ROOT).as_posix()}  ({path.stat().st_size:,} bytes)")


if __name__ == "__main__":
    main()
